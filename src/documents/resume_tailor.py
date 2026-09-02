from pathlib import Path

import docx
from ollama import Client
from sqlalchemy.orm import Session

from src.core.config import OLLAMA_HOST, OLLAMA_MODEL, TAILORED_RESUMES_DIR
from src.core.logger import logger
from src.database.models import Application, User, Vacancy
from src.documents.utils import looks_broken, safe_filename

_client = Client(host=OLLAMA_HOST)

HEADING_STYLES = ("Heading", "Title")

PROMPT_TEMPLATE = """You are tailoring a candidate's resume for a specific job posting.

Current resume (lines starting with "## " are section headings):
{resume_text}

They are applying for:
Title: {title}
Company: {company}
Job description:
{description}

Rewrite the resume to better match this job posting: emphasize the experience most relevant \
to it, reorder bullet points within each section so the most relevant ones come first, and use \
terminology/keywords from the job description wherever the candidate's real experience genuinely \
supports it.

Rules:
- Do NOT invent any experience, skill, employer, job title, degree, or achievement that is not \
already present in the original resume. Only rephrase, reorder, or emphasize existing content.
- Keep the same companies, dates, degrees, and other factual content unchanged.
- Keep the same output format: section headings on their own line prefixed with "## ", every \
other line as plain text.
- Never insert placeholder text or bracketed instructions like "[mention X]".

Return only the rewritten resume text, nothing else."""


def _skills_section_lines(text: str) -> list[str]:
    capture = False
    lines = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if line.startswith("## "):
            heading = line[3:].strip().lower()
            capture = any(k in heading for k in ("habilidade", "skill", "competenc", "tecnolog"))
            continue
        if capture and line:
            lines.append(line)
    return lines


def _find_fabricated_skills(original_text: str, tailored_text: str) -> list[str]:
    original_lower = original_text.lower()
    fabricated = []
    for line in _skills_section_lines(tailored_text):
        for item in line.split(","):
            item = item.strip()
            if item and item.lower() not in original_lower:
                fabricated.append(item)
    return fabricated


def is_supported_resume(path: str | None) -> bool:
    return bool(path) and Path(path).suffix.lower() == ".docx" and Path(path).exists()


def extract_resume_text(path: str) -> str:
    document = docx.Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith(HEADING_STYLES):
            lines.append(f"## {text}")
        else:
            lines.append(text)
    return "\n".join(lines)


def generate_tailored_resume_text(resume_text: str, vacancy: Vacancy) -> str:
    prompt = PROMPT_TEMPLATE.format(
        resume_text=resume_text,
        title=vacancy.title,
        company=vacancy.company,
        description=vacancy.description or "(no description provided)",
    )

    messages = [{"role": "user", "content": prompt}]
    for attempt in range(3):
        response = _client.chat(model=OLLAMA_MODEL, messages=messages, options={"temperature": 0.3})
        text = response["message"]["content"].strip()

        fabricated = _find_fabricated_skills(resume_text, text)
        if not looks_broken(text) and not fabricated:
            return text

        if fabricated:
            logger.warning(
                f"Tailored resume for vacancy {vacancy.id} invented skills on attempt {attempt + 1}: "
                f"{fabricated}, retrying"
            )
            feedback = (
                f"You added skills not present in the original resume: {', '.join(fabricated)}. "
                "Remove those and rewrite using only skills, technologies, and experience that "
                "already appear in the original resume."
            )
        else:
            logger.warning(
                f"Tailored resume for vacancy {vacancy.id} looked broken on attempt {attempt + 1}, retrying"
            )
            feedback = (
                "That contained placeholder brackets or garbled text. Rewrite it cleanly with no "
                "brackets and no garbled tokens, using only content already present in the "
                "original resume."
            )

        messages = [
            {"role": "user", "content": prompt},
            {"role": "assistant", "content": text},
            {"role": "user", "content": feedback},
        ]

    raise ValueError(f"Tailored resume generation kept producing invalid content for vacancy {vacancy.id}")


def render_resume_docx(text: str, vacancy: Vacancy) -> Path:
    filename = f"{vacancy.id}_{safe_filename(vacancy.company)}.docx"
    output_path = TAILORED_RESUMES_DIR / filename

    document = docx.Document()
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        else:
            document.add_paragraph(line)

    document.save(str(output_path))
    return output_path


def tailor_resume_for_application(session: Session, application: Application) -> Path | None:
    user = session.get(User, application.user_id)
    vacancy = session.get(Vacancy, application.vacancy_id)

    if not is_supported_resume(user.resume_path):
        logger.info(
            f"Skipping resume tailoring for application {application.id}: "
            f"resume_path is not a .docx file ({user.resume_path})"
        )
        return None

    resume_text = extract_resume_text(user.resume_path)
    tailored_text = generate_tailored_resume_text(resume_text, vacancy)
    output_path = render_resume_docx(tailored_text, vacancy)

    application.resume_used_path = str(output_path)
    session.commit()
    logger.info(f"Tailored resume generated for application {application.id}, vacancy {vacancy.id}")
    return output_path
